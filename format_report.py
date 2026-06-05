"""
format_report.py
Merges Sprint 1 + Sprint 2 markdown, parses it, and builds a fully
formatted DOCX using only python-docx (no pandoc required).

Usage:
    pip install python-docx
    python format_report.py
"""

import os
import sys
import re

PROJECT_PATH = os.path.dirname(os.path.abspath(__file__))
PROJECT_NAME = "StaySync_Salon_Reservation"

SPRINT1 = os.path.join(PROJECT_PATH, "PROJECT_REPORT_SPRINT1.md")
SPRINT2 = os.path.join(PROJECT_PATH, "PROJECT_REPORT_SPRINT2.md")
FINAL_DOCX = os.path.join(PROJECT_PATH, f"{PROJECT_NAME}_Report_Formatted.docx")


def step(msg):
    print(f"[STEP] {msg}")


def read_and_merge():
    step("Merging Sprint 1 + Sprint 2 markdown...")
    parts = []
    for fpath in [SPRINT1, SPRINT2]:
        if not os.path.exists(fpath):
            print(f"  ERROR: Not found: {fpath}")
            sys.exit(1)
        with open(fpath, "r", encoding="utf-8") as f:
            parts.append(f.read())
        print(f"  Read: {os.path.basename(fpath)}")

    combined = parts[0].rstrip() + "\n\n" + parts[1]
    # Remove the Sprint-2 title
    combined = re.sub(
        r"^#\s+SALON RESERVATION AND SCHEDULING SYSTEM.*Sprint 2\).*$",
        "",
        combined,
        flags=re.MULTILINE,
    )
    return combined


def clean(text):
    """Strip markdown artifacts from inline text."""
    if not text:
        return text
    text = text.replace("**", "")
    text = text.replace("__", "")
    text = text.replace("`", "")
    # leave single * for italic — just strip them too
    text = text.replace("*", "")
    return text


def parse_table_block(lines):
    """Parse a markdown table (list of lines) into rows of cells."""
    rows = []
    for line in lines:
        line = line.strip().strip("|")
        cells = [clean(c.strip()) for c in line.split("|")]
        rows.append(cells)
    # Remove separator row (the one with dashes)
    rows = [r for r in rows if not all(re.match(r"^[-:]+$", c) for c in r)]
    return rows


def build_docx(md_text):
    step("Building formatted DOCX from markdown...")

    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml
    except ImportError:
        print("  ERROR: python-docx not installed.")
        print("  Run:  pip install python-docx")
        sys.exit(1)

    doc = Document()

    # ---- Page setup ----
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1)

        # Footer with page number
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.clear()
        run = fp.add_run()
        fld = (
            '<w:fldSimple xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
            ' w:instr=" PAGE \\* MERGEFORMAT ">'
            '<w:r><w:rPr><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
            '<w:sz w:val="20"/></w:rPr><w:t>1</w:t></w:r></w:fldSimple>'
        )
        run._r.getparent().append(parse_xml(fld))

    # ---- Helpers ----
    def set_font(run, size=12, bold=False, name="Times New Roman"):
        run.font.name = name
        run.font.size = Pt(size)
        run.font.bold = bold
        rpr = run._r.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = parse_xml(
                '<w:rFonts {} w:ascii="{n}" w:hAnsi="{n}" '
                'w:eastAsia="{n}" w:cs="{n}"/>'.format(nsdecls("w"), n=name)
            )
            rpr.insert(0, rfonts)
        else:
            for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                rfonts.set(qn(attr), name)

    def set_spacing(para):
        """Set 1.5 line spacing."""
        pf = para.paragraph_format
        pf.line_spacing = Pt(18)  # 1.5 x 12pt

    def add_paragraph(text, size=12, bold=False, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
        p = doc.add_paragraph()
        run = p.add_run(clean(text))
        set_font(run, size=size, bold=bold)
        p.alignment = alignment
        set_spacing(p)
        return p

    def add_chapter_heading(text):
        # Page break before (skip for very first content)
        if len(doc.paragraphs) > 1:
            doc.add_page_break()
        p = doc.add_paragraph()
        run = p.add_run(clean(text))
        set_font(run, size=16, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p)
        return p

    def add_sub_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(clean(text))
        set_font(run, size=14, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p)
        return p

    def add_sub_sub_heading(text):
        p = doc.add_paragraph()
        run = p.add_run(clean(text))
        set_font(run, size=12, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_spacing(p)
        return p

    def add_table(rows):
        if not rows or len(rows) < 1:
            return
        ncols = max(len(r) for r in rows)
        # Pad rows to same number of cols
        for r in rows:
            while len(r) < ncols:
                r.append("")

        tbl = doc.add_table(rows=len(rows), cols=ncols)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Borders
        tbl_elem = tbl._tbl
        tbl_pr = tbl_elem.tblPr
        borders_xml = (
            '<w:tblBorders {}>'
            '<w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '<w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            '</w:tblBorders>'.format(nsdecls("w"))
        )
        old = tbl_pr.find(qn("w:tblBorders"))
        if old is not None:
            tbl_pr.remove(old)
        tbl_pr.append(parse_xml(borders_xml))

        for ri, row_data in enumerate(rows):
            row = tbl.rows[ri]
            for ci, cell_text in enumerate(row_data):
                cell = row.cells[ci]
                # Clear default empty paragraph
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                if ri == 0:
                    set_font(run, size=11, bold=True)
                    # Grey header background
                    shd = parse_xml(
                        '<w:shd {} w:fill="D9D9D9" w:val="clear"/>'.format(nsdecls("w"))
                    )
                    tc_pr = cell._tc.get_or_add_tcPr()
                    old_shd = tc_pr.find(qn("w:shd"))
                    if old_shd is not None:
                        tc_pr.remove(old_shd)
                    tc_pr.append(shd)
                else:
                    set_font(run, size=11, bold=False)

        doc.add_paragraph()  # spacing after table

    def add_mermaid_placeholder(lines):
        """Add mermaid code block as a placeholder text box."""
        p = doc.add_paragraph()
        run = p.add_run("[Diagram Placeholder - Render from Mermaid code below]")
        set_font(run, size=11, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add the mermaid code as a code block
        for line in lines:
            cp = doc.add_paragraph()
            cr = cp.add_run(line)
            set_font(cr, size=9, bold=False, name="Consolas")
            cp.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.add_paragraph()  # spacing

    # ---- Parse markdown line by line ----
    lines = md_text.split("\n")
    i = 0
    total = len(lines)
    is_first_heading = True

    while i < total:
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Skip horizontal rules
        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            i += 1
            continue

        # Mermaid code block
        if stripped.startswith("```mermaid"):
            i += 1
            mermaid_lines = []
            while i < total and not lines[i].strip().startswith("```"):
                mermaid_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            add_mermaid_placeholder(mermaid_lines)
            continue

        # Other code blocks — skip
        if stripped.startswith("```"):
            i += 1
            while i < total and not lines[i].strip().startswith("```"):
                i += 1
            i += 1  # skip closing ```
            continue

        # Table detection
        if "|" in stripped and i + 1 < total and re.match(r"^\s*\|?[\s\-:|]+\|", lines[i + 1]):
            table_lines = []
            while i < total and "|" in lines[i].strip():
                table_lines.append(lines[i])
                i += 1
            rows = parse_table_block(table_lines)
            if rows:
                add_table(rows)
            continue

        # Heading level 1: # CHAPTER ...
        if re.match(r"^#\s+", stripped) and not stripped.startswith("##"):
            text = re.sub(r"^#\s+", "", stripped)
            if is_first_heading:
                is_first_heading = False
            add_chapter_heading(text)
            i += 1
            continue

        # Heading level 2: ## CHAPTER ... or ## X.Y
        if re.match(r"^##\s+", stripped) and not stripped.startswith("###"):
            text = re.sub(r"^##\s+", "", stripped)
            # Check if it's a CHAPTER heading
            if re.match(r"CHAPTER\s+\d+", text, re.IGNORECASE):
                add_chapter_heading(text)
            else:
                add_sub_heading(text)
            i += 1
            continue

        # Heading level 3: ### X.Y.Z
        if re.match(r"^###\s+", stripped) and not stripped.startswith("####"):
            text = re.sub(r"^###\s+", "", stripped)
            # Check if it looks like X.Y.Z
            if re.match(r"\d+\.\d+\.\d+", text):
                add_sub_sub_heading(text)
            else:
                add_sub_sub_heading(text)
            i += 1
            continue

        # Heading level 4: ####
        if re.match(r"^#{4,}\s+", stripped):
            text = re.sub(r"^#{4,}\s+", "", stripped)
            add_sub_sub_heading(text)
            i += 1
            continue

        # Bullet / list items - collect consecutive ones
        if re.match(r"^[-*]\s+", stripped) or re.match(r"^\d+\.\s+", stripped):
            text = re.sub(r"^[-*]\s+", "", stripped)
            text = re.sub(r"^\d+\.\s+", "", text)
            p = doc.add_paragraph()
            run = p.add_run(clean(text))
            set_font(run, size=12)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_spacing(p)
            # Indent
            p.paragraph_format.left_indent = Inches(0.5)
            i += 1
            continue

        # Regular paragraph — collect continuation lines
        para_text = stripped
        i += 1
        while i < total:
            next_line = lines[i].strip()
            if not next_line:
                break
            if next_line.startswith("#"):
                break
            if next_line.startswith("```"):
                break
            if "|" in next_line and i + 1 < total and re.match(r"^\s*\|?[\s\-:|]+\|", lines[i + 1] if i + 1 < total else ""):
                break
            if re.match(r"^[-*]\s+", next_line):
                break
            if re.match(r"^-{3,}$", next_line):
                break
            para_text += " " + next_line
            i += 1

        # Skip if just markdown artifacts
        cleaned = clean(para_text).strip()
        if not cleaned:
            continue

        # Parenthetical / placeholder text in italics
        if cleaned.startswith("(") and cleaned.endswith(")"):
            p = doc.add_paragraph()
            run = p.add_run(cleaned)
            set_font(run, size=12)
            run.font.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            set_spacing(p)
            continue

        add_paragraph(para_text)

    # ---- Save ----
    doc.save(FINAL_DOCX)
    print(f"  Saved: {os.path.basename(FINAL_DOCX)}")


def main():
    print("=" * 60)
    print("  StaySync Report Formatter (No-Pandoc Edition)")
    print("=" * 60)
    print()

    md_text = read_and_merge()
    print()
    build_docx(md_text)

    print()
    print("=" * 60)
    print("  DONE!")
    print(f"  Output: {FINAL_DOCX}")
    print("=" * 60)


if __name__ == "__main__":
    main()
